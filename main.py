from pdfminer.high_level import extract_text
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage
import os
import inspect

cwd =os.path.dirname(os.path.abspath(inspect.getfile(inspect.currentframe())))

# List all the pdfs on the cwd
def list_pdfs_in_directory(directory):
    pdf_files = [f for f in os.listdir(directory) if f.lower().endswith('.pdf')]
    return pdf_files

available_pdfs = list_pdfs_in_directory(cwd)
print("Available PDFs in the directory:")
for idx, pdf in enumerate(available_pdfs, start=1):
    print(f"{idx}. {pdf}")

if available_pdfs:
    while True:
        try:
            selection = int(input("Select a PDF by number: "))
            if 1 <= selection <= len(available_pdfs):
                selected_pdf = available_pdfs[selection - 1]
                print(f"You selected: {selected_pdf}")
                break
            else:
                print("Invalid selection. Please enter a number from the list.")
        except ValueError:
            print("Please enter a valid number.")
else:
    print("No PDF files found in the directory.")

def extract_text_from_pdf(pdf_path):
    try:
        text = extract_text(pdf_path)
        return text
    except Exception as e:
        print(f"Error extracting text from {pdf_path}: {e}")
        return None

def get_pdf_metadata(pdf_path):
    num_pages = 0
    language = None
    with open(pdf_path, 'rb') as f:
        parser = PDFParser(f)
        doc = PDFDocument(parser)
        # Get number of pages
        try:
            num_pages = len(list(PDFPage.create_pages(doc)))
        except Exception:
            num_pages = 0
        # Try to get language from metadata
        try:
            if hasattr(doc, 'info') and doc.info:
                info = doc.info[0]
                if 'Lang' in info:
                    language = info['Lang']
                elif 'Language' in info:
                    language = info['Language']
                elif 'lang' in info:
                    language = info['lang']
                if isinstance(language, bytes):
                    language = language.decode(errors='ignore')
        except Exception:
            language = None
    return num_pages, language

pdf_path = os.path.join(cwd, selected_pdf)
extracted_text = extract_text_from_pdf(pdf_path)
num_pages, language = get_pdf_metadata(pdf_path)

print(f"\nPDF: {selected_pdf}")
print(f"Number of pages: {num_pages}")
print(f"Language: {language if language else 'Unknown'}\n")

if extracted_text:
    print(f"Extracted text from {selected_pdf}:\n")
    print(extracted_text)

def save_extracted_text(text, filename, file_format):
    if file_format == 'txt':
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(text)
    elif file_format == 'docx':
        from docx import Document
        import re

        def clean_text_for_docx(text):
            # Remove NULL bytes and problematic control characters
            return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)

        cleaned_text = clean_text_for_docx(text)
        doc = Document()
        doc.add_paragraph(cleaned_text)
        doc.save(filename)


def get_save_option():
    while True:
        print("\nHow would you like to save the extracted text?")
        print("1. Save as TXT file")
        print("2. Save as DOCX file")
        print("3. Do not save")
        choice = input("Enter your choice (1/2/3): ")
        if choice == '1':
            return 'txt'
        elif choice == '2':
            return 'docx'
        elif choice == '3':
            return None
        else:
            print("Invalid choice. Please try again.")


save_format = get_save_option()
if save_format:
    filename = f"{os.path.splitext(selected_pdf)[0]}_extracted.{save_format}"
    save_extracted_text(extracted_text, filename, save_format)
    print(f"Extracted text saved as {filename}")